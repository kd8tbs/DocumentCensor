import PyPDF2
from fpdf import FPDF
import re
from tkinter import*
from PyPDF2 import PdfReader
from PyPDF2 import PageObject, PdfReader
from docx import Document

def extract_text_from_doc(path):
    document = Document(path)
    text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
    return text


def save_text_as_doc(text, output_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)
   
def redact_string(input_string, regex_pattern):
    redacted_string = re.sub(regex_pattern, "[REDACTED]", input_string)
    return redacted_string

def shield(input_path, output_path):
    text = extract_text_from_doc(input_path)
    regex_arr = [
        r"\d\d\d-\d\d-\d\d\d\d",  # Social Security 123-12-1234
        r"\d\d/\d\d/\d\d\d\d",  # Birthdays 11/30/1980
        r"^[A-Za-z]{2}\d{2}-\d{5}-\d{4}$",  # Driver License Numbers D123456789123
        r"\d\d\d\d\d\d\d\d\d",  # Routing Number 123456789
        r"\(\d\d\d\)\d\d\d-\d\d\d\d"  # Phone Numbers (231)667-503
    ]

    for regex in regex_arr:
        text = redact_string(text, regex)

    save_text_as_doc(text, "output.pdf")
    



